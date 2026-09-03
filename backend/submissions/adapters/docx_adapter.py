from __future__ import annotations

from io import BytesIO

from docx import Document

from submissions.adapters.base import BaseSubmissionAdapter, ExtractedDocument
from submissions.text_sanitize import sanitize_json, sanitize_text


class DocxAdapter(BaseSubmissionAdapter):
    file_type = "docx"

    def extract(self, data: bytes, filename: str) -> ExtractedDocument:
        doc = Document(BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        return ExtractedDocument(
            text=sanitize_text(full_text),
            structure=sanitize_json({"paragraphs": paragraphs, "paragraph_count": len(paragraphs)}),
            source_ref=filename,
        )
